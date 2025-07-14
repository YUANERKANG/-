import pdfplumber as ppl
import docx


# 导入正则表达式相关库


def get_pdf_data(pdf_path="C:\\Users\\瑠衣\\Desktop\\软件杯\\智能简历解析\\简历pdf\\1.pdf"):
    """
    获取pdf的内容
    :param pdf_path: 路径
    """
    # 使用 PDFPlumber 库打开PDF文件
    pdf = ppl.open(pdf_path)

    TextList = []

    # 获得 PDFPlumber 的对象，可以查看其中的文字内容
    for page in pdf.pages:
        TextList.append(page.extract_text())
    text = ''.join(TextList)
    # print(text)

    return text


def get_word_data(word_path="C:\\Users\\瑠衣\\Desktop\\软件杯\\智能简历解析\\简历doc\\1.docx"):
    """
    获取word的内容
    :param word_path: 路径
    """
    # 使用 PDFPlumber 库打开word文件
    doc = docx.Document(word_path)

    print(doc.paragraphs)

    TextList = []
    # 读取自然段
    for paragraph in doc.paragraphs:
        TextList.append(paragraph.text)

    children = doc.element.body.iter()
    for child in children:
        # 通过类型判断目录
        if child.tag.endswith('txbx'):
            for ci in child.iter():
                if ci.tag.endswith('main}r'):
                    TextList.append(ci.text)
    # 读取表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                TextList.append(cell.text)

    text = ''.join(TextList)
    print(text)
    return text


def get_txt_data(txt_path="C:\\Users\\瑠衣\\Desktop\\软件杯\\智能简历解析\\简历txt\\3.txt"):
    f = open(txt_path, "r", encoding="utf-8")
    lines = f.readlines()  # 读取全部内容
    text = ''.join(lines)
    # print(text)
    return text



if __name__ == '__main__':
    text=get_txt_data("C:\\Users\\瑠衣\\Desktop\\软件杯\\智能简历解析\\简历txt\\17.txt")
    print(text)
    text=text.replace(' ', '')
    print(text)


